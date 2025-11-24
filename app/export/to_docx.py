from __future__ import annotations

from pathlib import Path
from typing import Dict

from docx import Document
from docx.shared import Inches

from app.models import ProcessingResult
from app.utils.image_io import imwrite


def export_docx(result: ProcessingResult, export_dir: Path, meta: Dict[str, str] | None = None) -> Path:
    """Create a DOCX report with metrics and annotated image."""
    export_dir.mkdir(parents=True, exist_ok=True)
    doc_path = export_dir / f"{result.target_id}.docx"
    meta = meta or {}

    document = Document()
    document.add_heading("Отчёт по мишени", level=1)
    if meta.get("organization"):
        document.add_paragraph(f"Организация: {meta['organization']}")
    if meta.get("author"):
        document.add_paragraph(f"Оператор: {meta['author']}")
    document.add_paragraph(f"ID: {result.target_id}")
    document.add_paragraph(f"Дата: {result.timestamp.strftime('%Y-%m-%d %H:%M:%S')}")

    metrics = result.metrics
    if metrics:
        summary = result.to_summary_dict()
        table = document.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Параметр"
        hdr_cells[1].text = "Значение"
        rows = [
            ("Количество", f"{summary.get('point_count', 0)}"),
            ("СТП X (мм)", f"{summary.get('mean_x_mm', 0.0):.2f}"),
            ("СТП Y (мм)", f"{summary.get('mean_y_mm', 0.0):.2f}"),
            ("Смещ. (мм)", f"{summary.get('displacement_mm', 0.0):.2f}"),
            ("Mean Radius", f"{summary.get('mean_radius_mm', 0.0):.2f}"),
            ("Extreme Spread", f"{summary.get('extreme_spread_mm', 0.0):.2f}"),
            ("R50", f"{summary.get('r50_mm', 0.0):.2f}"),
        ]
        for name, value in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = name
            row_cells[1].text = value

    overlay_path = _ensure_overlay_path(result, export_dir)
    if overlay_path:
        document.add_heading("Схема попаданий", level=2)
        document.add_picture(str(overlay_path), width=Inches(4.5))

    if result.points:
        document.add_heading("Попадания", level=2)
        points_table = document.add_table(rows=1, cols=4)
        hdr = points_table.rows[0].cells
        hdr[0].text = "#"
        hdr[1].text = "X (мм)"
        hdr[2].text = "Y (мм)"
        hdr[3].text = "Источник"
        for point in sorted(result.points, key=lambda p: p.id):
            row = points_table.add_row().cells
            row[0].text = str(point.id)
            row[1].text = f"{point.x_mm:.2f}"
            row[2].text = f"{point.y_mm:.2f}"
            row[3].text = point.source

    document.save(str(doc_path))
    return doc_path


def _ensure_overlay_path(result: ProcessingResult, export_dir: Path) -> str | None:
    if result.overlay_path and Path(result.overlay_path).exists():
        return result.overlay_path
    if result.overlay_image is None:
        return None
    fallback = export_dir / f"{result.target_id}_overlay.png"
    if imwrite(fallback, result.overlay_image):
        result.overlay_path = str(fallback)
    else:
        raise RuntimeError(f"Не удалось сохранить изображение оверлея для DOCX: {fallback}")
    return result.overlay_path


__all__ = ["export_docx"]
